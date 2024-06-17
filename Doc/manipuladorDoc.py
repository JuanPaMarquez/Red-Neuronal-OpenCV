import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx2pdf import convert
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from reportlab.platypus import Paragraph, SimpleDocTemplate

# Obtener el directorio actual y el directorio del proyecto
current_dir = os.path.dirname(os.path.realpath(__file__))
project_dir = os.path.dirname(current_dir)

# Ruta del archivo .docx original y del archivo modificado
input_data_path = os.path.join(project_dir, 'fga152_asistencia_clase.docx')
output_data_path = os.path.join(project_dir, 'fga152_asistencia_clase_modificado.docx')
pdf_output_path = os.path.join(project_dir, 'fga152_asistencia_clase_modificado.pdf')

def create_pdf(output_path, fecha, grupo, aula, docente, curso, tema, sede):
    # Crear un lienzo (canvas) para el PDF
    c = canvas.Canvas(output_path, pagesize=letter)
    
    # Definir estilos para los campos
    style = ParagraphStyle(name='Normal', fontName='Helvetica', fontSize=12)
    centered_style = ParagraphStyle(name='Centered', parent=style, alignment=TA_CENTER)

    # Definir las coordenadas para cada campo
    fields = [
        (fecha, 100, 700),
        (f"GRUPO: {grupo}", 100, 680),
        (f"AULA: {aula}", 100, 660),
        (f"DOCENTE: {docente}", 100, 640),
        (f"CURSO: {curso}", 100, 620),
        (f"TEMA: {tema}", 100, 600),
        (f"SEDE: {sede}", 100, 580)
    ]

    # Agregar cada campo al lienzo en las coordenadas especificadas
    for text, x, y in fields:
        p = Paragraph(text, style)
        p.wrapOn(c, 200, 20)
        p.drawOn(c, x, y)
    
    # Guardar el lienzo y cerrar el archivo PDF
    c.save()

def modify_document(doc_path, output_path):
    # Abrir el documento original
    doc = Document(doc_path)
    
    # Buscar el texto 'FECHA:   _____________________________' y reemplazarlo
    for paragraph in doc.paragraphs:
        if 'FECHA:   _____________________________' in paragraph.text:
            for run in paragraph.runs:
                if 'FECHA:   _____________________________' in run.text:
                    run.text = run.text.replace('FECHA:   _____________________________', 'FECHA:   _16 de junio de 2024_')
    
    # Guardar el documento modificado
    doc.save(output_path)

def insert_data_into_table(doc_path, output_path, data, attended_docs):
    # Abrir el documento original
    doc = Document(doc_path)
    
    # Asumir que la tabla es la primera en el documento
    table = doc.tables[0]
    
    # Limpiar las filas existentes (excepto la fila de encabezado) y luego insertar datos
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)
    
    # Insertar los datos en la tabla, comenzando desde la segunda fila
    for idx, (doc_num, name, program) in enumerate(data):
        row_cells = table.add_row().cells
        
        # Establecer y formatear el número en la primera celda
        number_paragraph = row_cells[0].paragraphs[0]
        number_paragraph.clear()  # Limpiar el contenido existente
        number_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinear el párrafo al centro
        number_run = number_paragraph.add_run()
        number_run.bold = True
        number_run.font.size = Pt(11)
        number_run.font.name = 'Arial'
        number_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        number_run.text = str(idx + 1)
        
        # Insertar y formatear los datos en las demás celdas
        row_cells[1].text = name
        row_cells[2].text = str(doc_num)
        row_cells[3].text = program
        
        # Formatear la celda de firma con centrado y texto correspondiente
        signature_paragraph = row_cells[4].paragraphs[0]
        signature_paragraph.clear()  # Limpiar el contenido existente
        signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alinear el párrafo al centro
        signature_run = signature_paragraph.add_run()
        signature_run.font.size = Pt(11)
        signature_run.font.name = 'Arial'
        signature_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        signature_run.text = 'ASISTIÓ' if doc_num in attended_docs else 'NO ASISTIÓ'
        
        # Aplicar el formato Arial 11 sin negrita a cada celda de la fila, excepto la primera columna
        for cell in row_cells[1:4]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    
    # Guardar el documento modificado
    doc.save(output_path)

def convert_docx_to_pdf(input_path, output_path):
    # Convertir el documento .docx a .pdf usando docx2pdf
    convert(input_path, output_path)

# Ejemplo de uso
if __name__ == "__main__":

    fecha = "16 de junio de 2024"
    grupo = "A"
    aula = "101"
    docente = "Profesor Ejemplo"
    curso = "Programación Avanzada"
    tema = "Introducción a reportlab"
    sede = "Campus Principal"
    
    # Crear el PDF con los campos especificados
    # Datos de ejemplo
    data = [(12346546, 'Raiul', 'Ing Sistemas'), (456789, 'Paul', 'Ing Sistemas')]
    attended_docs = [12346546, 456789]  # Documentos de los estudiantes que asistieron
    
    # Modificar el documento .docx
    modify_document(input_data_path, output_data_path)
    
    # Insertar datos en la tabla del documento modificado
    insert_data_into_table(output_data_path, output_data_path, data, attended_docs)
    
    # Convertir el documento modificado a PDF
    convert_docx_to_pdf(output_data_path, pdf_output_path)

    #create_pdf(pdf_output_path, fecha, grupo, aula, docente, curso, tema, sede)

    print(f"Documento PDF guardado en: {pdf_output_path}")
